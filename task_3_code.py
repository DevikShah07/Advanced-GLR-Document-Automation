import streamlit as st
import pandas as pd
import numpy as np
import json
import re
import logging
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
from pathlib import Path
import io
import base64
from concurrent.futures import ThreadPoolExecutor
import asyncio

# Document Processing
import pdfplumber
from docx import Document
from docx2pdf import convert
import pytesseract
import easyocr
from PIL import Image
from pdf2image import convert_from_bytes

# AI/ML
from sentence_transformers import SentenceTransformer
import torch
from sklearn.metrics.pairwise import cosine_similarity

# Data Validation
from pydantic import BaseModel, Field, validator
from email_validator import validate_email, EmailNotValidError
import phonenumbers
import dateparser

# Visualization
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots

# API Integration
import openai
import requests
from dotenv import load_dotenv
import os

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# =====================================
# CONFIGURATION AND CONSTANTS
# =====================================

class Config:
    """Application configuration"""
    OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
    OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"
    
    # Model configurations
    EMBEDDING_MODEL = "all-MiniLM-L6-v2"
    BACKUP_EMBEDDING_MODEL = "all-mpnet-base-v2"
    LLM_MODEL = "mistralai/mistral-7b-instruct"
    
    # Thresholds
    CONFIDENCE_THRESHOLD_HIGH = 0.85
    CONFIDENCE_THRESHOLD_MEDIUM = 0.65
    CONFIDENCE_THRESHOLD_LOW = 0.45
    
    # Processing
    MAX_TEXT_CHUNK_SIZE = 5000
    MAX_CONCURRENT_EXTRACTIONS = 5

# =====================================
# DATA MODELS AND VALIDATION
# =====================================

class GLRField(BaseModel):
    """Individual field with validation and confidence scoring"""
    key: str
    value: str
    confidence: float = Field(ge=0.0, le=1.0)
    source: str = Field(description="extraction_method")
    validated: bool = False
    
    @validator('value')
    def validate_value(cls, v, values):
        """Apply field-specific validation"""
        key = values.get('key', '')
        
        if key == 'XM8_CONTACT_EMAIL' and v != "NOT_FOUND":
            try:
                validate_email(v)
            except EmailNotValidError:
                raise ValueError(f"Invalid email format: {v}")
        
        if key == 'XM8_CONTACT_PHONE' and v != "NOT_FOUND":
            try:
                parsed = phonenumbers.parse(v, "US")
                if not phonenumbers.is_valid_number(parsed):
                    raise ValueError(f"Invalid phone number: {v}")
            except phonenumbers.NumberParseException:
                raise ValueError(f"Invalid phone number format: {v}")
        
        return v

class GLRDocument(BaseModel):
    """Complete GLR document with all fields and metadata"""
    fields: Dict[str, GLRField]
    overall_confidence: float
    extraction_timestamp: datetime
    processing_time: float
    generated_content: Dict[str, str] = {}
    
    class Config:
        arbitrary_types_allowed = True

# =====================================
# ADVANCED FIELD VALIDATORS
# =====================================

class FieldValidators:
    """Advanced field validation with regex patterns and business rules"""
    
    PATTERNS = {
        'XM8_CLAIM_NUM': [
            r'[A-Z]{2,4}[-_]?\d{6,12}',
            r'\d{8,15}',
            r'CLM[-_]?\d{6,10}'
        ],
        'XM8_POLICY_NUM': [
            r'[A-Z]{2,4}[-_]?\d{6,15}',
            r'POL[-_]?\d{6,12}',
            r'\d{8,20}'
        ],
        'XM8_DATE_OF_LOSS': [
            r'\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}',
            r'\d{4}[\/\-]\d{1,2}[\/\-]\d{1,2}'
        ],
        'XM8_CONTACT_PHONE': [
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        ],
        'XM8_CONTACT_EMAIL': [
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        ],
        'XM8_ESTIMATED_DAMAGE': [
            r'\$[\d,]+\.?\d*',
            r'[\d,]+\.?\d*\s*dollars?'
        ]
    }
    
    @classmethod
    def validate_field(cls, field_key: str, value: str) -> Tuple[bool, str, float]:
        """Validate a field using regex patterns and return confidence score"""
        if value in ["NOT_FOUND", "", None]:
            return False, value, 0.0
        
        patterns = cls.PATTERNS.get(field_key, [])
        if not patterns:
            return True, value, 0.7  # Default confidence for fields without patterns
        
        for pattern in patterns:
            matches = re.findall(pattern, value, re.IGNORECASE)
            if matches:
                best_match = max(matches, key=len) if isinstance(matches[0], str) else matches[0]
                confidence = min(1.0, len(best_match) / 10)  # Simple confidence based on match length
                return True, best_match, confidence
        
        return False, value, 0.3  # Low confidence for non-matching patterns

    @classmethod
    def normalize_date(cls, date_str: str) -> str:
        """Normalize date to MM/DD/YYYY format"""
        if date_str in ["NOT_FOUND", "", None]:
            return date_str
        
        try:
            parsed_date = dateparser.parse(date_str)
            if parsed_date:
                return parsed_date.strftime("%m/%d/%Y")
        except:
            pass
        
        return date_str

    @classmethod
    def normalize_currency(cls, amount_str: str) -> str:
        """Normalize currency to $X,XXX.XX format"""
        if amount_str in ["NOT_FOUND", "", None]:
            return amount_str
        
        # Extract numeric value
        numeric_str = re.sub(r'[^\d.]', '', amount_str)
        try:
            amount = float(numeric_str)
            return f"${amount:,.2f}"
        except ValueError:
            return amount_str

# =====================================
# ADVANCED TEXT EXTRACTION
# =====================================

class AdvancedTextExtractor:
    """Multi-modal text extraction with OCR and advanced parsing"""
    
    def __init__(self):
        self.ocr_reader = easyocr.Reader(['en'])
        
    def extract_from_pdf(self, pdf_bytes: bytes) -> Dict[str, Any]:
        """Extract text using multiple methods"""
        results = {
            'text_chunks': [],
            'ocr_chunks': [],
            'tables': [],
            'metadata': {}
        }
        
        try:
            # Method 1: pdfplumber for text and tables
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        results['text_chunks'].append({
                            'page': page_num + 1,
                            'text': text.strip(),
                            'method': 'pdfplumber'
                        })
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        results['tables'].append({
                            'page': page_num + 1,
                            'table': table,
                            'method': 'pdfplumber'
                        })
            
            # Method 2: OCR for images and poor-quality text
            try:
                images = convert_from_bytes(pdf_bytes, dpi=300)
                for page_num, image in enumerate(images):
                    ocr_text = self.ocr_reader.readtext(np.array(image))
                    if ocr_text:
                        combined_text = ' '.join([item[1] for item in ocr_text if item[2] > 0.5])
                        if combined_text.strip():
                            results['ocr_chunks'].append({
                                'page': page_num + 1,
                                'text': combined_text.strip(),
                                'method': 'ocr',
                                'confidence': np.mean([item[2] for item in ocr_text])
                            })
            except Exception as e:
                logger.warning(f"OCR extraction failed: {e}")
        
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            
        return results

    def combine_extractions(self, extraction_results: Dict[str, Any]) -> List[str]:
        """Combine and deduplicate text from multiple extraction methods"""
        all_texts = []
        
        # Add pdfplumber text
        for chunk in extraction_results['text_chunks']:
            all_texts.append(chunk['text'])
        
        # Add OCR text (if significantly different from pdfplumber)
        for chunk in extraction_results['ocr_chunks']:
            ocr_text = chunk['text']
            # Simple deduplication: check if OCR text is significantly different
            is_duplicate = any(
                self._similarity_score(ocr_text, existing) > 0.8 
                for existing in all_texts
            )
            if not is_duplicate:
                all_texts.append(ocr_text)
        
        # Add table text
        for table_data in extraction_results['tables']:
            table_text = self._table_to_text(table_data['table'])
            all_texts.append(table_text)
        
        return all_texts

    def _similarity_score(self, text1: str, text2: str) -> float:
        """Calculate similarity between two texts"""
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = len(words1.intersection(words2))
        union = len(words1.union(words2))
        
        return intersection / union if union > 0 else 0.0

    def _table_to_text(self, table: List[List[str]]) -> str:
        """Convert table data to searchable text"""
        if not table:
            return ""
        
        text_parts = []
        for row in table:
            if row:
                cleaned_row = [str(cell).strip() for cell in row if cell]
                text_parts.append(" ".join(cleaned_row))
        
        return "\n".join(text_parts)

# =====================================
# ADVANCED FIELD EXTRACTION ENGINE
# =====================================

class AdvancedFieldExtractor:
    """Multi-method field extraction with ensemble scoring"""
    
    def __init__(self):
        self.primary_model = SentenceTransformer(Config.EMBEDDING_MODEL)
        self.backup_model = SentenceTransformer(Config.BACKUP_EMBEDDING_MODEL)
        self.validators = FieldValidators()
        
        # Field definitions with semantic descriptions
        self.field_definitions = {
            'XM8_CLAIM_NUM': {
                'description': 'claim number file number reference number case number',
                'aliases': ['claim #', 'claim number', 'file #', 'file number', 'reference', 'case number'],
                'weight': 1.0
            },
            'XM8_INSURED_NAME': {
                'description': 'insured name policyholder name customer name client name',
                'aliases': ['insured', 'policyholder', 'customer', 'client', 'name of insured'],
                'weight': 1.0
            },
            'XM8_POLICY_NUM': {
                'description': 'policy number insurance policy contract number',
                'aliases': ['policy #', 'policy number', 'contract #', 'insurance policy'],
                'weight': 1.0
            },
            'XM8_DATE_OF_LOSS': {
                'description': 'date of loss incident date damage date occurrence date',
                'aliases': ['date of loss', 'incident date', 'loss date', 'damage date', 'occurrence'],
                'weight': 1.0
            },
            'XM8_LOCATION': {
                'description': 'property address location site address damage location',
                'aliases': ['address', 'location', 'property address', 'site', 'damage location'],
                'weight': 0.9
            },
            'XM8_ADJUSTER_NAME': {
                'description': 'adjuster name inspector name examiner name',
                'aliases': ['adjuster', 'inspector', 'examiner', 'adjuster name'],
                'weight': 0.9
            },
            'XM8_CONTACT_PHONE': {
                'description': 'phone number contact number telephone mobile number',
                'aliases': ['phone', 'telephone', 'mobile', 'contact number', 'phone number'],
                'weight': 0.8
            },
            'XM8_CONTACT_EMAIL': {
                'description': 'email address contact email electronic mail',
                'aliases': ['email', 'e-mail', 'contact email', 'email address'],
                'weight': 0.8
            },
            'XM8_CAUSE_OF_LOSS': {
                'description': 'cause of loss damage cause peril reason for damage',
                'aliases': ['cause', 'cause of loss', 'peril', 'damage cause', 'reason'],
                'weight': 0.9
            },
            'XM8_ESTIMATED_DAMAGE': {
                'description': 'estimated damage cost repair estimate damage amount',
                'aliases': ['estimate', 'damage cost', 'repair cost', 'estimated damage'],
                'weight': 0.9
            },
            'XM8_DEDUCTIBLE': {
                'description': 'deductible amount policy deductible insurance deductible',
                'aliases': ['deductible', 'policy deductible', 'insurance deductible'],
                'weight': 0.8
            },
            'XM8_INSPECTION_DATE': {
                'description': 'inspection date survey date examination date',
                'aliases': ['inspection date', 'survey date', 'examination date', 'inspected'],
                'weight': 0.8
            },
            'XM8_PROPERTY_TYPE': {
                'description': 'property type building type structure type',
                'aliases': ['property type', 'building type', 'structure', 'type of property'],
                'weight': 0.7
            },
            'XM8_COVERAGE_TYPE': {
                'description': 'coverage type insurance coverage policy coverage',
                'aliases': ['coverage', 'insurance coverage', 'policy coverage', 'coverage type'],
                'weight': 0.7
            },
            'XM8_ADDITIONAL_NOTES': {
                'description': 'additional notes comments remarks observations',
                'aliases': ['notes', 'comments', 'remarks', 'observations', 'additional'],
                'weight': 0.6
            }
        }

    def extract_fields_ensemble(self, text_chunks: List[str]) -> Dict[str, GLRField]:
        """Extract fields using ensemble of methods"""
        results = {}
        
        # Method 1: Embedding-based extraction
        embedding_results = self._extract_with_embeddings(text_chunks)
        
        # Method 2: Regex pattern matching
        regex_results = self._extract_with_regex(text_chunks)
        
        # Method 3: Keyword proximity search
        keyword_results = self._extract_with_keywords(text_chunks)
        
        # Combine results with weighted scoring
        for field_key in self.field_definitions.keys():
            candidates = []
            
            # Collect candidates from all methods
            if field_key in embedding_results:
                candidates.append(embedding_results[field_key])
            if field_key in regex_results:
                candidates.append(regex_results[field_key])
            if field_key in keyword_results:
                candidates.append(keyword_results[field_key])
            
            # Select best candidate
            if candidates:
                best_candidate = max(candidates, key=lambda x: x.confidence)
                results[field_key] = best_candidate
            else:
                results[field_key] = GLRField(
                    key=field_key,
                    value="NOT_FOUND",
                    confidence=0.0,
                    source="no_extraction"
                )
        
        return results

    def _extract_with_embeddings(self, text_chunks: List[str]) -> Dict[str, GLRField]:
        """Extract using semantic embeddings"""
        results = {}
        
        try:
            # Combine text chunks
            combined_text = " ".join(text_chunks[:10])  # Limit for performance
            
            # Split into sentences
            sentences = re.split(r'[.!?]+', combined_text)
            sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
            
            if not sentences:
                return results
            
            # Generate embeddings
            sentence_embeddings = self.primary_model.encode(sentences)
            
            for field_key, field_def in self.field_definitions.items():
                query_embedding = self.primary_model.encode([field_def['description']])
                
                # Calculate similarities
                similarities = cosine_similarity(query_embedding, sentence_embeddings)[0]
                
                # Find best match
                best_idx = np.argmax(similarities)
                best_score = similarities[best_idx]
                
                if best_score > Config.CONFIDENCE_THRESHOLD_LOW:
                    # Extract potential value from best sentence
                    best_sentence = sentences[best_idx]
                    extracted_value = self._extract_value_from_sentence(best_sentence, field_key)
                    
                    # Validate extracted value
                    is_valid, normalized_value, validation_confidence = self.validators.validate_field(
                        field_key, extracted_value
                    )
                    
                    # Combine confidences
                    final_confidence = (best_score * 0.7) + (validation_confidence * 0.3)
                    
                    results[field_key] = GLRField(
                        key=field_key,
                        value=normalized_value,
                        confidence=final_confidence,
                        source="embeddings",
                        validated=is_valid
                    )
        
        except Exception as e:
            logger.error(f"Embedding extraction failed: {e}")
        
        return results

    def _extract_with_regex(self, text_chunks: List[str]) -> Dict[str, GLRField]:
        """Extract using regex patterns"""
        results = {}
        combined_text = " ".join(text_chunks)
        
        for field_key in self.field_definitions.keys():
            patterns = self.validators.PATTERNS.get(field_key, [])
            
            for pattern in patterns:
                matches = re.findall(pattern, combined_text, re.IGNORECASE)
                if matches:
                    best_match = max(matches, key=len) if isinstance(matches[0], str) else matches[0]
                    
                    # Validate match
                    is_valid, normalized_value, confidence = self.validators.validate_field(
                        field_key, best_match
                    )
                    
                    if confidence > 0.5:
                        results[field_key] = GLRField(
                            key=field_key,
                            value=normalized_value,
                            confidence=confidence,
                            source="regex",
                            validated=is_valid
                        )
                        break
        
        return results

    def _extract_with_keywords(self, text_chunks: List[str]) -> Dict[str, GLRField]:
        """Extract using keyword proximity search"""
        results = {}
        combined_text = " ".join(text_chunks).lower()
        
        for field_key, field_def in self.field_definitions.items():
            aliases = field_def['aliases']
            
            for alias in aliases:
                # Find keyword positions
                keyword_positions = [m.start() for m in re.finditer(re.escape(alias.lower()), combined_text)]
                
                for pos in keyword_positions:
                    # Extract surrounding context
                    start = max(0, pos - 50)
                    end = min(len(combined_text), pos + 200)
                    context = combined_text[start:end]
                    
                    # Try to extract value from context
                    extracted_value = self._extract_value_from_context(context, field_key, alias)
                    
                    if extracted_value and extracted_value != "NOT_FOUND":
                        # Validate extracted value
                        is_valid, normalized_value, confidence = self.validators.validate_field(
                            field_key, extracted_value
                        )
                        
                        if confidence > 0.4:
                            results[field_key] = GLRField(
                                key=field_key,
                                value=normalized_value,
                                confidence=confidence * 0.8,  # Slightly lower confidence for keyword method
                                source="keywords",
                                validated=is_valid
                            )
                            break
                
                if field_key in results:
                    break
        
        return results

    def _extract_value_from_sentence(self, sentence: str, field_key: str) -> str:
        """Extract specific value from a sentence based on field type"""
        sentence = sentence.strip()
        
        if field_key in ['XM8_CLAIM_NUM', 'XM8_POLICY_NUM']:
            # Look for alphanumeric patterns
            patterns = [r'[A-Z]{2,4}[-_]?\d{6,15}', r'\d{8,20}']
            for pattern in patterns:
                matches = re.findall(pattern, sentence)
                if matches:
                    return matches[0]
        
        elif field_key in ['XM8_DATE_OF_LOSS', 'XM8_INSPECTION_DATE']:
            # Look for date patterns
            date_patterns = [
                r'\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}',
                r'\d{4}[\/\-]\d{1,2}[\/\-]\d{1,2}',
                r'[A-Za-z]{3,9}\s+\d{1,2},?\s+\d{4}'
            ]
            for pattern in date_patterns:
                matches = re.findall(pattern, sentence)
                if matches:
                    return matches[0]
        
        elif field_key in ['XM8_ESTIMATED_DAMAGE', 'XM8_DEDUCTIBLE']:
            # Look for currency amounts
            currency_patterns = [r'\$[\d,]+\.?\d*', r'[\d,]+\.?\d*\s*dollars?']
            for pattern in currency_patterns:
                matches = re.findall(pattern, sentence, re.IGNORECASE)
                if matches:
                    return matches[0]
        
        elif field_key == 'XM8_CONTACT_EMAIL':
            # Look for email addresses
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            matches = re.findall(email_pattern, sentence)
            if matches:
                return matches[0]
        
        elif field_key == 'XM8_CONTACT_PHONE':
            # Look for phone numbers
            phone_patterns = [
                r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
                r'\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
            ]
            for pattern in phone_patterns:
                matches = re.findall(pattern, sentence)
                if matches:
                    return matches[0]
        
        # For other fields, return the sentence (to be processed by LLM)
        return sentence[:100]  # Limit length

    def _extract_value_from_context(self, context: str, field_key: str, keyword: str) -> str:
        """Extract value from keyword context"""
        # Split context around keyword
        parts = context.split(keyword.lower())
        if len(parts) < 2:
            return "NOT_FOUND"
        
        # Look for value after keyword
        after_keyword = parts[1].strip()
        
        # Common separators
        separators = [':', '=', '-', '–', '—']
        for sep in separators:
            if sep in after_keyword:
                potential_value = after_keyword.split(sep)[1].strip()
                # Extract first meaningful part
                words = potential_value.split()
                if words:
                    if field_key in ['XM8_CLAIM_NUM', 'XM8_POLICY_NUM']:
                        return words[0]
                    elif field_key in ['XM8_INSURED_NAME', 'XM8_ADJUSTER_NAME']:
                        return ' '.join(words[:3])  # Assume names are 1-3 words
                    else:
                        return ' '.join(words[:5])  # General case
        
        # If no separator, take first few words after keyword
        words = after_keyword.split()
        if words:
            if field_key in ['XM8_CLAIM_NUM', 'XM8_POLICY_NUM']:
                return words[0]
            elif field_key in ['XM8_INSURED_NAME', 'XM8_ADJUSTER_NAME']:
                return ' '.join(words[:3])
            else:
                return ' '.join(words[:5])
        
        return "NOT_FOUND"

# =====================================
# LLM INTEGRATION
# =====================================

class LLMExtractor:
    """Advanced LLM integration with structured output"""
    
    def __init__(self):
        self.api_key = Config.OPENROUTER_API_KEY
        self.base_url = Config.OPENROUTER_BASE_URL
        
    def extract_with_llm(self, text_chunks: List[str], missing_fields: List[str] = None) -> Dict[str, GLRField]:
        """Extract fields using LLM with structured output"""
        if not self.api_key:
            logger.error("OpenRouter API key not configured")
            return {}
        
        # Combine text chunks
        combined_text = " ".join(text_chunks[:5])  # Limit for API
        
        # Create targeted prompt
        prompt = self._create_extraction_prompt(combined_text, missing_fields)
        
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": Config.LLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.1,
                "max_tokens": 2000
            }
            
            response = requests.post(
                f"{self.base_url}/chat/completions",
                headers=headers,
                json=payload,
                timeout=30
            )
            
            if response.status_code == 200:
                response_data = response.json()
                content = response_data["choices"][0]["message"]["content"]
                
                # Parse JSON response
                return self._parse_llm_response(content)
            else:
                logger.error(f"LLM API error: {response.status_code} - {response.text}")
                return {}
                
        except Exception as e:
            logger.error(f"LLM extraction failed: {e}")
            return {}

    def _create_extraction_prompt(self, text: str, missing_fields: List[str] = None) -> str:
        """Create optimized extraction prompt"""
        
        fields_to_extract = missing_fields or [
            "XM8_CLAIM_NUM", "XM8_INSURED_NAME", "XM8_POLICY_NUM", "XM8_DATE_OF_LOSS",
            "XM8_LOCATION", "XM8_ADJUSTER_NAME", "XM8_CONTACT_PHONE", "XM8_CONTACT_EMAIL",
            "XM8_CAUSE_OF_LOSS", "XM8_ESTIMATED_DAMAGE", "XM8_DEDUCTIBLE", "XM8_INSPECTION_DATE",
            "XM8_PROPERTY_TYPE", "XM8_COVERAGE_TYPE", "XM8_ADDITIONAL_NOTES"
        ]
        
        json_structure = {field: "string" for field in fields_to_extract}
        
        return f"""
You are an expert insurance document processor. Extract the following information from the provided text and return ONLY a valid JSON object with the exact keys shown below.

**CRITICAL INSTRUCTIONS:**
- Return ONLY valid JSON format - no other text, explanations, or markdown
- Use the exact key names provided (including XM8_ prefixes)
- If a field cannot be found, use "NOT_FOUND" as the value
- For dates, use MM/DD/YYYY format
- For monetary amounts, include currency symbol and use format $X,XXX.XX
- For names, use proper capitalization

**TEXT TO ANALYZE:**
{text[:4000]}

**REQUIRED JSON OUTPUT FORMAT:**
{json.dumps(json_structure, indent=2)}

**FIELD EXTRACTION GUIDELINES:**
- CLAIM_NUM: Look for claim numbers, file numbers, or reference numbers
- INSURED_NAME: Individual or business name of the insured party
- POLICY_NUM: Insurance policy number or contract number
- DATE_OF_LOSS: When the incident/damage occurred
- LOCATION: Property address or location of loss
- ADJUSTER_NAME: Insurance adjuster or inspector name
- CONTACT_PHONE: Phone number for primary contact
- CONTACT_EMAIL: Email address for primary contact  
- CAUSE_OF_LOSS: What caused the damage (fire, water, wind, etc.)
- ESTIMATED_DAMAGE: Total estimated repair/replacement cost
- DEDUCTIBLE: Policy deductible amount
- INSPECTION_DATE: When property was inspected
- PROPERTY_TYPE: Type of property (residential, commercial, etc.)
- COVERAGE_TYPE: Type of insurance coverage
- ADDITIONAL_NOTES: Any other relevant information

Return only the JSON object - no additional text.
"""

    def _parse_llm_response(self, content: str) -> Dict[str, GLRField]:
        """Parse LLM JSON response into GLRField objects"""
        try:
            # Clean response
            content = content.strip()
            if content.startswith('```json'):
                content = content.replace('```json', '').replace('```', '')
            
            # Parse JSON
            data = json.loads(content)
            
            results = {}
            for key, value in data.items():
                if key.startswith('XM8_'):
                    # Apply field-specific normalization
                    if key in ['XM8_DATE_OF_LOSS', 'XM8_INSPECTION_DATE']:
                        value = FieldValidators.normalize_date(value)
                    elif key in ['XM8_ESTIMATED_DAMAGE', 'XM8_DEDUCTIBLE']:
                        value = FieldValidators.normalize_currency(value)
                    
                    # Validate field
                    is_valid, normalized_value, confidence = FieldValidators.validate_field(key, value)
                    
                    # Adjust confidence based on LLM source
                    final_confidence = confidence * 0.9 if value != "NOT_FOUND" else 0.0
                    
                    results[key] = GLRField(
                        key=key,
                        value=normalized_value,
                        confidence=final_confidence,
                        source="llm",
                        validated=is_valid
                    )
            
            return results
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse LLM JSON response: {e}")
            return {}
        except Exception as e:
            logger.error(f"Error processing LLM response: {e}")
            return {}

# =====================================
# CONTENT GENERATION
# =====================================

class ContentGenerator:
    """Generate professional narrative content for GLR sections"""
    
    def __init__(self):
        self.api_key = Config.OPENROUTER_API_KEY
        self.base_url = Config.OPENROUTER_BASE_URL
    
    def generate_narrative_sections(self, extracted_fields: Dict[str, GLRField], text_chunks: List[str]) -> Dict[str, str]:
        """Generate professional narrative sections"""
        if not self.api_key:
            return {}
        
        sections = {}
        
        try:
            # Generate Inspection Summary
            sections['INSPECTION_SUMMARY'] = self._generate_inspection_summary(extracted_fields, text_chunks)
            
            # Generate Risk Assessment
            sections['RISK_ASSESSMENT'] = self._generate_risk_assessment(extracted_fields, text_chunks)
            
            # Generate Recommendations
            sections['RECOMMENDATIONS'] = self._generate_recommendations(extracted_fields, text_chunks)
            
        except Exception as e:
            logger.error(f"Content generation failed: {e}")
        
        return sections
    
    def _generate_inspection_summary(self, fields: Dict[str, GLRField], text_chunks: List[str]) -> str:
        """Generate professional inspection summary"""
        context = self._build_context(fields, text_chunks)
        
        prompt = f"""
Write a professional 2-3 paragraph inspection summary for an insurance GLR report based on the following information:

Context: {context}

Requirements:
- Professional insurance industry tone
- Factual and objective
- Include key details about the inspection
- 150-250 words
- Do not include placeholder text

Write only the summary paragraphs, no headers or formatting.
"""
        
        return self._call_llm_for_content(prompt)
    
    def _generate_risk_assessment(self, fields: Dict[str, GLRField], text_chunks: List[str]) -> str:
        """Generate risk assessment section"""
        context = self._build_context(fields, text_chunks)
        
        prompt = f"""
Write a professional risk assessment for an insurance GLR report based on the following information:

Context: {context}

Requirements:
- Assess potential risks and exposures
- Professional insurance industry language
- Include mitigation recommendations if applicable
- 100-200 words
- Objective and analytical tone

Write only the risk assessment content, no headers.
"""
        
        return self._call_llm_for_content(prompt)
    
    def _generate_recommendations(self, fields: Dict[str, GLRField], text_chunks: List[str]) -> str:
        """Generate recommendations section"""
        context = self._build_context(fields, text_chunks)
        
        prompt = f"""
Write professional recommendations for an insurance GLR report based on the following information:

Context: {context}

Requirements:
- Specific, actionable recommendations
- Professional insurance industry tone
- Focus on claim handling and next steps
- 100-150 words
- Use bullet points if appropriate

Write only the recommendations content, no headers.
"""
        
        return self._call_llm_for_content(prompt)
    
    def _build_context(self, fields: Dict[str, GLRField], text_chunks: List[str]) -> str:
        """Build context from extracted fields and text"""
        context_parts = []
        
        # Add high-confidence field values
        for field_key, field in fields.items():
            if field.confidence > 0.6 and field.value != "NOT_FOUND":
                clean_key = field_key.replace('XM8_', '').replace('_', ' ').title()
                context_parts.append(f"{clean_key}: {field.value}")
        
        # Add relevant text excerpts
        combined_text = " ".join(text_chunks[:3])  # Limit for context
        if combined_text:
            context_parts.append(f"Additional context: {combined_text[:500]}...")
        
        return " | ".join(context_parts)
    
    def _call_llm_for_content(self, prompt: str) -> str:
        """Call LLM API for content generation"""
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": Config.LLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 500
            }
            
            response = requests.post(
                f"{self.base_url}/chat/completions",
                headers=headers,
                json=payload,
                timeout=30
            )
            
            if response.status_code == 200:
                response_data = response.json()
                return response_data["choices"][0]["message"]["content"].strip()
            else:
                logger.error(f"Content generation API error: {response.status_code}")
                return "Content generation unavailable."
                
        except Exception as e:
            logger.error(f"Content generation failed: {e}")
            return "Content generation unavailable."

# =====================================
# MAIN PROCESSING ENGINE
# =====================================

class GLRProcessingEngine:
    """Main processing engine that orchestrates all components"""
    
    def __init__(self):
        self.text_extractor = AdvancedTextExtractor()
        self.field_extractor = AdvancedFieldExtractor()
        self.llm_extractor = LLMExtractor()
        self.content_generator = ContentGenerator()
        
    def process_documents(self, template_bytes: bytes, pdf_files: List[bytes]) -> GLRDocument:
        """Process GLR documents with full pipeline"""
        start_time = datetime.now()
        
        try:
            # Step 1: Extract text from all PDFs
            st.info("🔍 Extracting text from PDF reports...")
            all_text_chunks = []
            
            for i, pdf_bytes in enumerate(pdf_files):
                st.write(f"Processing PDF {i+1}/{len(pdf_files)}")
                extraction_results = self.text_extractor.extract_from_pdf(pdf_bytes)
                text_chunks = self.text_extractor.combine_extractions(extraction_results)
                all_text_chunks.extend(text_chunks)
            
            if not all_text_chunks:
                raise ValueError("No text could be extracted from PDF files")
            
            # Step 2: Extract fields using ensemble methods
            st.info("🧠 Extracting fields using AI ensemble...")
            extracted_fields = self.field_extractor.extract_fields_ensemble(all_text_chunks)
            
            # Step 3: LLM fallback for low-confidence fields
            st.info("🤖 Enhancing extraction with LLM...")
            low_confidence_fields = [
                key for key, field in extracted_fields.items() 
                if field.confidence < Config.CONFIDENCE_THRESHOLD_MEDIUM
            ]
            
            if low_confidence_fields:
                llm_results = self.llm_extractor.extract_with_llm(all_text_chunks, low_confidence_fields)
                
                # Merge LLM results (prefer LLM for low-confidence fields)
                for field_key, llm_field in llm_results.items():
                    if field_key in extracted_fields:
                        current_field = extracted_fields[field_key]
                        if llm_field.confidence > current_field.confidence:
                            extracted_fields[field_key] = llm_field
            
            # Step 4: Generate narrative content
            st.info("📝 Generating professional narrative sections...")
            generated_content = self.content_generator.generate_narrative_sections(
                extracted_fields, all_text_chunks
            )
            
            # Step 5: Calculate overall confidence
            confidences = [field.confidence for field in extracted_fields.values()]
            overall_confidence = np.mean(confidences) if confidences else 0.0
            
            # Step 6: Create final document
            processing_time = (datetime.now() - start_time).total_seconds()
            
            glr_document = GLRDocument(
                fields=extracted_fields,
                overall_confidence=overall_confidence,
                extraction_timestamp=datetime.now(),
                processing_time=processing_time,
                generated_content=generated_content
            )
            
            return glr_document
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            raise

# =====================================
# DOCUMENT PROCESSING
# =====================================

class DocumentProcessor:
    """Handle DOCX template processing and PDF conversion"""
    
    @staticmethod
    def fill_template(template_bytes: bytes, glr_document: GLRDocument) -> bytes:
        """Fill DOCX template with extracted data"""
        try:
            # Ensure we have valid bytes
            if not template_bytes:
                raise ValueError("Template bytes are empty")
            
            if not isinstance(template_bytes, bytes):
                raise ValueError("Template data must be bytes")
            
            # Load template from bytes
            doc = Document(io.BytesIO(template_bytes))
            
            # Replace field placeholders in paragraphs
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                for field_key, field in glr_document.fields.items():
                    placeholder = f"[{field_key}]"
                    if placeholder in original_text:
                        # Replace the placeholder with the field value
                        new_text = original_text.replace(placeholder, field.value)
                        paragraph.text = new_text
                        original_text = new_text
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        original_text = cell.text
                        for field_key, field in glr_document.fields.items():
                            placeholder = f"[{field_key}]"
                            if placeholder in original_text:
                                new_text = original_text.replace(placeholder, field.value)
                                cell.text = new_text
                                original_text = new_text
            
            # Add generated content to specific sections
            for section_key, content in glr_document.generated_content.items():
                placeholder = f"[{section_key}]"
                for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, content)
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            return doc_bytes.getvalue()
            
        except Exception as e:
            logger.error(f"Template filling failed: {e}")
            logger.error(f"Template bytes type: {type(template_bytes)}")
            logger.error(f"Template bytes length: {len(template_bytes) if template_bytes else 'None'}")
            raise

# =====================================
# VISUALIZATION COMPONENTS
# =====================================

class VisualizationComponents:
    """Create interactive visualizations for the Streamlit app"""
    
    @staticmethod
    def create_confidence_chart(glr_document: GLRDocument) -> go.Figure:
        """Create confidence score visualization"""
        field_names = []
        confidences = []
        colors = []
        
        for field_key, field in glr_document.fields.items():
            clean_name = field_key.replace('XM8_', '').replace('_', ' ').title()
            field_names.append(clean_name)
            confidences.append(field.confidence)
            
            # Color coding based on confidence
            if field.confidence >= Config.CONFIDENCE_THRESHOLD_HIGH:
                colors.append('green')
            elif field.confidence >= Config.CONFIDENCE_THRESHOLD_MEDIUM:
                colors.append('orange')
            else:
                colors.append('red')
        
        fig = go.Figure(data=[
            go.Bar(
                x=field_names,
                y=confidences,
                marker_color=colors,
                text=[f"{c:.1%}" for c in confidences],
                textposition='auto',
            )
        ])
        
        fig.update_layout(
            title="Field Extraction Confidence Scores",
            xaxis_title="Fields",
            yaxis_title="Confidence Score",
            yaxis=dict(range=[0, 1]),
            height=400
        )
        
        return fig
    
    @staticmethod
    def create_extraction_methods_pie(glr_document: GLRDocument) -> go.Figure:
        """Create pie chart of extraction methods used"""
        method_counts = {}
        
        for field in glr_document.fields.values():
            method = field.source
            method_counts[method] = method_counts.get(method, 0) + 1
        
        fig = go.Figure(data=[
            go.Pie(
                labels=list(method_counts.keys()),
                values=list(method_counts.values()),
                hole=0.3
            )
        ])
        
        fig.update_layout(
            title="Extraction Methods Distribution",
            height=300
        )
        
        return fig

# =====================================
# STREAMLIT APPLICATION
# =====================================

def main():
    """Main Streamlit application"""
    st.set_page_config(
        page_title="Advanced GLR Document Automation",
        page_icon="📋",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    st.title("🚀 Advanced GLR Document Automation System")
    st.markdown("---")
    
    # Initialize session state
    if 'processing_complete' not in st.session_state:
        st.session_state.processing_complete = False
    if 'glr_document' not in st.session_state:
        st.session_state.glr_document = None
    if 'filled_docx' not in st.session_state:
        st.session_state.filled_docx = None
    if 'template_bytes' not in st.session_state:  # Add this line
        st.session_state.template_bytes = None
    
    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Configuration
        with st.expander("🔧 API Settings"):
            api_key = st.text_input(
                "OpenRouter API Key",
                type="password",
                value=os.getenv("OPENROUTER_API_KEY", ""),
                help="Required for LLM fallback and content generation"
            )
            if api_key:
                os.environ["OPENROUTER_API_KEY"] = api_key
                Config.OPENROUTER_API_KEY = api_key
        
        # Processing Settings
        with st.expander("🎛️ Processing Settings"):
            confidence_threshold = st.slider(
                "Confidence Threshold",
                min_value=0.0,
                max_value=1.0,
                value=0.65,
                step=0.05,
                help="Fields below this threshold will use LLM fallback"
            )
            Config.CONFIDENCE_THRESHOLD_MEDIUM = confidence_threshold
            
            enable_ocr = st.checkbox("Enable OCR", value=True, help="Extract text from images")
            enable_content_generation = st.checkbox("Generate Narrative Content", value=True)
    
    # Main interface
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📄 Upload Documents")
        
        # Template upload
        template_file = st.file_uploader(
            "Upload DOCX Template",
            type=['docx'],
            help="Upload your GLR template with [XM8_...] placeholders"
        )
        
        # PDF reports upload
        pdf_files = st.file_uploader(
            "Upload PDF Reports",
            type=['pdf'],
            accept_multiple_files=True,
            help="Upload one or more PDF inspection reports"
        )
        
        # Process button
        if st.button("🚀 Process Documents", type="primary", disabled=not (template_file and pdf_files)):
            if not Config.OPENROUTER_API_KEY:
                st.error("⚠️ Please configure your OpenRouter API key in the sidebar")
            else:
                with st.spinner("Processing documents..."):
                    try:
                        # Initialize processing engine
                        engine = GLRProcessingEngine()
                
                        # Read template bytes ONCE and store
                        template_bytes = template_file.read()
                
                        # Process documents
                        pdf_bytes_list = [pdf.read() for pdf in pdf_files]
                        glr_document = engine.process_documents(template_bytes, pdf_bytes_list)
                
                        # Fill template using the same bytes
                        processor = DocumentProcessor()
                        filled_docx = processor.fill_template(template_bytes, glr_document)
                
                        # Store in session state
                        st.session_state.glr_document = glr_document
                        st.session_state.filled_docx = filled_docx
                        st.session_state.template_bytes = template_bytes  # Store for later use
                        st.session_state.processing_complete = True
                
                        st.success("✅ Processing completed successfully!")
                
                    except Exception as e:
                        st.error(f"❌ Processing failed: {str(e)}")
                        logger.error(f"Processing error: {e}")
    
    with col2:
        st.header("📊 Processing Results")
        
        if st.session_state.processing_complete and st.session_state.glr_document:
            glr_doc = st.session_state.glr_document
            
            # Overall metrics
            col_a, col_b, col_c = st.columns(3)
            with col_a:
                st.metric("Overall Confidence", f"{glr_doc.overall_confidence:.1%}")
            with col_b:
                st.metric("Processing Time", f"{glr_doc.processing_time:.1f}s")
            with col_c:
                fields_found = sum(1 for f in glr_doc.fields.values() if f.value != "NOT_FOUND")
                st.metric("Fields Found", f"{fields_found}/{len(glr_doc.fields)}")
            
            # Confidence visualization
            vis = VisualizationComponents()
            confidence_chart = vis.create_confidence_chart(glr_doc)
            st.plotly_chart(confidence_chart, use_container_width=True)
    
    # Field editing interface
    if st.session_state.processing_complete and st.session_state.glr_document:
        st.header("✏️ Review & Edit Extracted Fields")
        
        glr_doc = st.session_state.glr_document
        
        # Create editable form
        with st.form("field_editor"):
            edited_fields = {}
            
            # Group fields by confidence for better UX
            high_conf_fields = {k: v for k, v in glr_doc.fields.items() if v.confidence >= Config.CONFIDENCE_THRESHOLD_HIGH}
            med_conf_fields = {k: v for k, v in glr_doc.fields.items() if Config.CONFIDENCE_THRESHOLD_MEDIUM <= v.confidence < Config.CONFIDENCE_THRESHOLD_HIGH}
            low_conf_fields = {k: v for k, v in glr_doc.fields.items() if v.confidence < Config.CONFIDENCE_THRESHOLD_MEDIUM}
            
            # High confidence fields
            if high_conf_fields:
                st.subheader("🟢 High Confidence Fields")
                for field_key, field in high_conf_fields.items():
                    clean_name = field_key.replace('XM8_', '').replace('_', ' ').title()
                    edited_fields[field_key] = st.text_input(
                        f"{clean_name} ({field.confidence:.1%})",
                        value=field.value,
                        key=f"edit_{field_key}"
                    )
            
            # Medium confidence fields
            if med_conf_fields:
                st.subheader("🟡 Medium Confidence Fields")
                for field_key, field in med_conf_fields.items():
                    clean_name = field_key.replace('XM8_', '').replace('_', ' ').title()
                    edited_fields[field_key] = st.text_input(
                        f"{clean_name} ({field.confidence:.1%})",
                        value=field.value,
                        key=f"edit_{field_key}"
                    )
            
            # Low confidence fields
            if low_conf_fields:
                st.subheader("🔴 Low Confidence Fields - Please Review")
                for field_key, field in low_conf_fields.items():
                    clean_name = field_key.replace('XM8_', '').replace('_', ' ').title()
                    edited_fields[field_key] = st.text_input(
                        f"{clean_name} ({field.confidence:.1%})",
                        value=field.value,
                        key=f"edit_{field_key}"
                    )
            
            # Update button
            if st.form_submit_button("🔄 Update Fields", type="primary"):
                # Update the GLR document with edited values
                for field_key, new_value in edited_fields.items():
                    if field_key in glr_doc.fields:
                        glr_doc.fields[field_key].value = new_value
                        glr_doc.fields[field_key].confidence = 1.0  # User-confirmed
                        glr_doc.fields[field_key].source = "user_edit"
    
                # Regenerate filled document using stored template bytes
                try:
                    processor = DocumentProcessor()
                    # Use stored template bytes instead of re-reading the file
                    template_bytes = st.session_state.get('template_bytes')
                    if template_bytes:
                        filled_docx = processor.fill_template(template_bytes, glr_doc)
                        st.session_state.filled_docx = filled_docx
                        st.success("✅ Fields updated successfully!")
                    else:
                        st.error("❌ Template data not found. Please re-upload and process.")
                except Exception as e:
                    st.error(f"❌ Failed to update document: {str(e)}")
    
    # Download section
    if st.session_state.processing_complete and st.session_state.filled_docx:
        st.header("⬇️ Download Results")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # Download DOCX
            st.download_button(
                label="📄 Download DOCX",
                data=st.session_state.filled_docx,
                file_name=f"GLR_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        with col2:
            # Download PDF (converted)
            try:
                pdf_bytes = DocumentProcessor.convert_to_pdf(st.session_state.filled_docx)
                st.download_button(
                    label="📑 Download PDF",
                    data=pdf_bytes,
                    file_name=f"GLR_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf"
                )
            except Exception as e:
                st.error(f"PDF conversion failed: {str(e)}")
        
        with col3:
            # Download extraction report
            if st.session_state.glr_document:
                report_data = {
                    "extraction_summary": {
                        "timestamp": st.session_state.glr_document.extraction_timestamp.isoformat(),
                        "processing_time": st.session_state.glr_document.processing_time,
                        "overall_confidence": st.session_state.glr_document.overall_confidence
                    },
                    "fields": {
                        k: {
                            "value": v.value,
                            "confidence": v.confidence,
                            "source": v.source,
                            "validated": v.validated
                        }
                        for k, v in st.session_state.glr_document.fields.items()
                    }
                }
                
                st.download_button(
                    label="📊 Download Report",
                    data=json.dumps(report_data, indent=2),
                    file_name=f"Extraction_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json"
                )
    
    # Analytics dashboard
    if st.session_state.processing_complete and st.session_state.glr_document:
        st.header("📈 Processing Analytics")
        
        glr_doc = st.session_state.glr_document
        vis = VisualizationComponents()
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Extraction methods pie chart
            methods_chart = vis.create_extraction_methods_pie(glr_doc)
            st.plotly_chart(methods_chart, use_container_width=True)
        
        with col2:
            # Field validation status
            validation_data = {
                'Validated': sum(1 for f in glr_doc.fields.values() if f.validated),
                'Not Validated': sum(1 for f in glr_doc.fields.values() if not f.validated)
            }
            
            fig = go.Figure(data=[
                go.Pie(
                    labels=list(validation_data.keys()),
                    values=list(validation_data.values()),
                    hole=0.3
                )
            ])
            fig.update_layout(title="Field Validation Status", height=300)
            st.plotly_chart(fig, use_container_width=True)

if __name__ == "__main__":
    main()